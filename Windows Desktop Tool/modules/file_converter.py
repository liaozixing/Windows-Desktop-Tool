import os
import sys
from PyQt5.QtGui import QImage, QPainter
from PyQt5.QtSvg import QSvgRenderer
from PyQt5.QtCore import QSize, Qt
from PIL import Image
import pandas as pd

# 动态导入处理
try:
    from pdf2docx import Converter as PDFConverter
except ImportError:
    PDFConverter = None

try:
    from docx2pdf import convert as docx_to_pdf_convert
except ImportError:
    docx_to_pdf_convert = None

try:
    from docx import Document
except ImportError:
    Document = None

def svg_to_ico(svg_path, ico_path):
    """ SVG 转 ICO """
    try:
        renderer = QSvgRenderer(svg_path)
        if not renderer.isValid():
            return False, "无效的 SVG 文件"
        
        # 创建 256x256 高清图
        image = QImage(256, 256, QImage.Format_ARGB32)
        image.fill(Qt.transparent)
        painter = QPainter(image)
        renderer.render(painter)
        painter.end()
        
        if image.save(ico_path, "ICO"):
            return True, "转换成功"
        return False, "保存 ICO 失败"
    except Exception as e:
        return False, str(e)

def image_convert(input_path, output_path, output_format):
    """ 通用图片格式转换 (PNG, JPG, BMP, WebP, etc.) """
    try:
        img = Image.open(input_path)
        # 如果是转 JPG，需要去掉透明通道
        if output_format.upper() in ["JPG", "JPEG"] and img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.save(output_path, output_format.upper())
        return True, "转换成功"
    except Exception as e:
        return False, str(e)

def pdf_to_word(pdf_path, docx_path):
    """ PDF 转 Word """
    if PDFConverter is None:
        return False, "缺少依赖库 pdf2docx"
    try:
        cv = PDFConverter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return True, "转换成功"
    except Exception as e:
        return False, str(e)

def word_to_pdf(docx_path, pdf_path):
    """ Word 转 PDF (需要系统安装有 MS Word) """
    if docx_to_pdf_convert is None:
        return False, "缺少依赖库 docx2pdf"
    try:
        docx_to_pdf_convert(docx_path, pdf_path)
        return True, "转换成功"
    except Exception as e:
        return False, str(e)

def word_to_excel(docx_path, excel_path):
    """ 提取 Word 中的表格到 Excel """
    if Document is None:
        return False, "缺少依赖库 python-docx"
    try:
        doc = Document(docx_path)
        all_tables = []
        for i, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            df = pd.DataFrame(data)
            all_tables.append(df)
        
        if not all_tables:
            return False, "Word 文档中未找到任何表格"
            
        with pd.ExcelWriter(excel_path) as writer:
            for i, df in enumerate(all_tables):
                df.to_excel(writer, sheet_name=f'Table_{i+1}', index=False, header=False)
        return True, "成功提取表格到 Excel"
    except Exception as e:
        return False, str(e)

def excel_to_word(excel_path, docx_path):
    """ 将 Excel 工作表转为 Word 表格 """
    if Document is None:
        return False, "缺少依赖库 python-docx"
    try:
        xls = pd.ExcelFile(excel_path)
        doc = Document()
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            doc.add_heading(f'Sheet: {sheet_name}', level=1)
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            # 添加表头
            for j in range(df.shape[1]):
                table.cell(0, j).text = str(df.columns[j])
            # 添加数据
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    table.cell(i + 1, j).text = str(df.values[i, j])
            doc.add_page_break()
            
        doc.save(docx_path)
        return True, "成功将 Excel 转为 Word 表格"
    except Exception as e:
        return False, str(e)
